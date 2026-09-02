from flask import Flask, render_template, request, jsonify, send_file
import google.generativeai as genai
import os
import PyPDF2
import docx
import markdown
import io
from weasyprint import HTML

app = Flask(__name__)

# إعداد مفتاح Gemini باستخدام المكتبة المستقرة
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

# تعريف النموذج المستقر
model = genai.GenerativeModel('gemini-3.5-flash-lite') if GEMINI_API_KEY else None

user_usage = {}
FREE_LIMIT = 3

def extract_text_from_pdf(file_stream):
    reader = PyPDF2.PdfReader(file_stream)
    text = ""
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text

def extract_text_from_docx(file_stream):
    doc = docx.Document(file_stream)
    text = "\n".join([p.text for p in doc.paragraphs if p.text])
    return text

@app.route('/')
def home():
    return render_template(
        'index.html',
        supabase_url=os.environ.get("SUPABASE_URL", ""),
        supabase_key=os.environ.get("SUPABASE_ANON_KEY", "")
    )

# مسار robots.txt البرمجي لحل مشكلة أدوات مشرفي الموقع نهائياً
@app.route('/robots.txt')
def robots_txt():
    return "User-agent: *\nAllow: /", 200, {'Content-Type': 'text/plain; charset=utf-8'}

# مسار ads.txt البرمجي الخاص بـ Google AdSense
@app.route('/ads.txt')
def ads_txt():
    return "google.com, pub-7946557086083356, DIRECT, f08c47fec0942fa0", 200, {'Content-Type': 'text/plain; charset=utf-8'}

@app.route('/optimize', methods=['POST'])
def optimize_cv():
    user_ip = request.remote_addr
    current_usage = user_usage.get(user_ip, 0)

    if current_usage >= FREE_LIMIT:
        return jsonify({
            'result': 'لقد استهلكت محاولاتك الـ 3 المجانية لهذا اليوم. / You have consumed your 3 free trials for today.'
        }), 403

    try:
        if not model:
            return jsonify({'result': 'مفتاح GEMINI_API_KEY غير معرف في البيئة. / GEMINI_API_KEY is not set.'}), 500

        job_title = request.form.get('job_title', '')
        language = request.form.get('language', 'English')
        cv_text = request.form.get('cv_text', '')

        if 'cv_file' in request.files and request.files['cv_file'].filename != '':
            file = request.files['cv_file']
            filename = file.filename.lower()
            if filename.endswith('.pdf'):
                cv_text = extract_text_from_pdf(file.stream)
            elif filename.endswith('.docx'):
                cv_text = extract_text_from_docx(file.stream)

        if not cv_text.strip():
            return jsonify({'result': 'يرجى إدخال نص أو رفع ملف سيرة ذاتية صالح. / Please enter text or upload a valid CV file.'}), 400

        prompt = f"""
You are an expert ATS Resume Optimizer.
Optimize the following CV text specifically for the position of '{job_title}'.
Target Output Language: {language}.

Requirements:
1. Make it fully ATS-friendly with standard headers and keywords.
2. Structure the content with Markdown headings (##), bold text, and bullet points.
3. Highlight relevant technical skills, tools, and achievements.

CV Content:
{cv_text}
"""
        response = model.generate_content(prompt)
        raw_text = response.text
        formatted_html = markdown.markdown(raw_text)

        user_usage[user_ip] = current_usage + 1

        return jsonify({
            'raw_text': raw_text,
            'formatted_html': formatted_html,
            'remaining_tries': FREE_LIMIT - user_usage[user_ip]
        })

    except Exception as e:
        return jsonify({'result': f"حدث خطأ أثناء معالجة الطلب: {str(e)}"}), 500

@app.route('/download/word', methods=['POST'])
def download_word():
    content = request.form.get('content', '')
    doc = docx.Document()
    doc.add_heading('Optimized CV - AI ATS', 0)
    for line in content.split('\n'):
        doc.add_paragraph(line)
        
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name='Optimized_Resume.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/download/pdf', methods=['POST'])
def download_pdf():
    content = request.form.get('content', '')
    direction = request.form.get('direction', 'ltr') # استقبال اتجاه الصفحة من الواجهة
    lang_attr = 'ar' if direction == 'rtl' else 'en'
    
    html_content = f"""
    <!DOCTYPE html>
    <html lang="{lang_attr}" dir="{direction}">
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: DejaVu Sans, Arial, sans-serif; padding: 20px; color: #111; line-height: 1.6; direction: {direction}; text-align: {'right' if direction == 'rtl' else 'left'}; }}
            h1, h2, h3 {{ color: #2563eb; }}
            hr {{ border: 0; border-top: 1px solid #ccc; margin: 15px 0; }}
        </style>
    </head>
    <body>
        <h2>{"السيرة الذاتية المحسنة (ATS Optimized)" if direction == 'rtl' else "Optimized Resume (ATS Optimized)"}</h2>
        <hr>
        <div>{content}</div>
    </body>
    </html>
    """
    
    pdf_bytes = HTML(string=html_content).write_pdf()
    file_stream = io.BytesIO(pdf_bytes)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name='Optimized_Resume.pdf',
        mimetype='application/pdf'
    )

if __name__ == '__main__':
    app.run(debug=True)